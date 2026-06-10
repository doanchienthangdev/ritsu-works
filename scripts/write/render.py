#!/usr/bin/env python3
"""render.py — render ONE /write output format from a markdown file.

  python render.py <markdown_file> <fmt> --out=PATH [--title=..] [--author=..]
                   [--style=ritsu] [--repo-root=DIR]

fmt: html | pdf (WeasyPrint) | docx (python-docx). `md`/`default` are handled in node.
Clean CONTENT layout (not the translation book-framing of scripts/translate/render.py).
Reuses scripts/translate/design.py load_tokens() for the design-system palette when present;
falls back to the Ritsu brand palette otherwise. Emits one line of JSON: {ok, out, fmt, warnings, error?}.
"""
import sys, os, re, json, html as H
from pathlib import Path

FALLBACK = {  # Ritsu "Electric Cyan V2" (knowledge/design-systems.yaml ritsu)
    "accent": "#0ABCD0", "ink": "#0B1620", "muted": "#5B6B76",
    "bg": "#FFFFFF", "rule": "#E4EAEE",
    "serif": "Georgia, 'Source Serif 4', 'Times New Roman', serif",
    "sans": "'Inter', -apple-system, 'Segoe UI', Helvetica, Arial, sans-serif",
}

def opts(argv):
    o = {}
    for a in argv:
        if a.startswith("--"):
            k, _, v = a[2:].partition("=")
            o[k] = v if v != "" else True
    return o

def load_palette(style, repo_root):
    warns = []
    pal = dict(FALLBACK)
    try:
        sys.path.insert(0, str(Path(repo_root) / "scripts" / "translate"))
        import design  # noqa
        tokens, _ = design.load_tokens(style, repo_root)
        # design.load_tokens shape varies; pull common keys defensively.
        if isinstance(tokens, dict):
            # NOTE muted: prefer a *foreground* token (readable grey for h3/byline/blockquote text).
            # design.load_tokens emits both `mutedForeground` (readable) and `muted` (a LIGHT surface
            # tint, e.g. claude #F1ECE2). Picking `muted` first made muted-colored text near-invisible
            # on light pages; prefer the foreground tokens and keep `muted`/`secondary` as last resort.
            cmap = {"accent": ["accent", "primary", "brand"], "ink": ["ink", "text", "foreground"],
                    "bg": ["bg", "background"],
                    "muted": ["mutedForeground", "muted-foreground", "muted-soft", "secondary", "muted"],
                    "rule": ["rule", "border", "hairline"]}
            for k, cands in cmap.items():
                for c in cands:
                    if tokens.get(c):
                        pal[k] = tokens[c]; break
    except Exception as e:  # noqa
        warns.append(f"palette: using fallback Ritsu palette ({type(e).__name__})")
    return pal, warns

def md_to_html_body(text):
    import markdown
    return markdown.markdown(text, extensions=["extra", "sane_lists", "smarty", "tables"])

def page_css(p):
    return f"""
    @page {{ size: A4; margin: 22mm 20mm; }}
    body {{ font-family: {p['serif']}; color: {p['ink']}; line-height: 1.62; font-size: 11.5pt; }}
    h1, h2, h3 {{ font-family: {p['sans']}; color: {p['ink']}; line-height: 1.2; font-weight: 700; }}
    h1 {{ font-size: 26pt; margin: 0 0 4pt; border-bottom: 3px solid {p['accent']}; padding-bottom: 8pt; }}
    h2 {{ font-size: 15pt; margin: 22pt 0 6pt; }}
    h3 {{ font-size: 12.5pt; margin: 16pt 0 4pt; color: {p['muted']}; }}
    p {{ margin: 0 0 9pt; }}
    a {{ color: {p['accent']}; }}
    blockquote {{ border-left: 3px solid {p['accent']}; margin: 10pt 0; padding: 2pt 0 2pt 14pt; color: {p['muted']}; font-style: italic; }}
    code {{ font-family: 'SF Mono', Menlo, monospace; background: #F2F5F7; padding: 1px 4px; border-radius: 3px; font-size: 10pt; }}
    pre {{ background: #F2F5F7; padding: 10pt; border-radius: 6px; overflow-x: auto; }}
    img {{ max-width: 100%; }}
    table {{ border-collapse: collapse; width: 100%; margin: 10pt 0; }}
    th, td {{ border: 1px solid {p['rule']}; padding: 6pt 8pt; text-align: left; font-size: 10.5pt; }}
    th {{ background: {p['accent']}11; }}
    .byline {{ font-family: {p['sans']}; color: {p['muted']}; font-size: 10pt; margin: 2pt 0 18pt; }}
    """

def render_html(md, title, author, pal):
    body = md_to_html_body(md)
    byline = f'<div class="byline">{H.escape(author)}</div>' if author else ""
    head = f'<h1>{H.escape(title)}</h1>{byline}' if title else ""
    return f"""<!doctype html><html><head><meta charset="utf-8"><title>{H.escape(title or "")}</title>
<style>{page_css(pal)}</style></head><body>{head}{body}</body></html>"""

def main():
    argv = sys.argv[1:]
    src = next((a for a in argv if not a.startswith("--")), None)
    o = opts(argv)
    fmt = None
    nonflag = [a for a in argv if not a.startswith("--")]
    if len(nonflag) >= 2:
        src, fmt = nonflag[0], nonflag[1]
    out = o.get("out"); title = o.get("title") or ""; author = o.get("author") or ""
    style = o.get("style") or "ritsu"; repo_root = o.get("repo-root") or os.getcwd()
    warns = []
    try:
        md = Path(src).read_text(encoding="utf-8")
        # Drop leading HTML comment(s) — the /write orchestrator records the framework
        # choice as an HTML comment on line 1 (write/orchestrator §3.5). It's scaffolding,
        # never output, and (being line 1) it otherwise defeats the title H1-strip below.
        md = re.sub(r"^\s*(?:<!--.*?-->\s*)+", "", md, flags=re.S)
        # strip a leading H1 (we render title separately for pdf/html/docx)
        if title:
            md = re.sub(r"^#\s+.*\n", "", md.lstrip(), count=1)
        pal, pw = load_palette(style, repo_root); warns += pw

        if fmt == "html":
            Path(out).write_text(render_html(md, title, author, pal), encoding="utf-8")
        elif fmt == "pdf":
            # The weasyprint module is broken on the anaconda interpreter (missing native
            # libs) AND in the standalone CLI here, but the HOMEBREW python has a working
            # weasyprint (that's the interpreter scripts/translate uses for PDF). So:
            # render the HTML on this (anaconda) python, then shell the homebrew weasyprint.
            import subprocess, tempfile
            html_str = render_html(md, title, author, pal)
            with tempfile.NamedTemporaryFile("w", suffix=".html", delete=False, encoding="utf-8") as fh:
                fh.write(html_str); tmp = fh.name
            weasy_py = os.environ.get("WRITE_PY_WEASY") or os.environ.get("TRANSLATE_PY_WEASY") or "/opt/homebrew/bin/python3"
            # base_url = the source markdown's directory so relative embeds
            # (enrich images ![](01.png) + dataviz charts ![](chart.svg)) resolve;
            # the temp HTML lives in /tmp, so without this they'd 404 in the PDF.
            base = os.path.dirname(os.path.abspath(src))
            code = ("from weasyprint import HTML; "
                    f"HTML(filename={tmp!r}, base_url={base!r}).write_pdf({out!r})")
            r = subprocess.run([weasy_py, "-c", code], capture_output=True, text=True)
            try: os.unlink(tmp)
            except OSError: pass
            if r.returncode != 0:
                print(json.dumps({"ok": False, "fmt": fmt, "error": f"weasyprint(homebrew) failed: {(r.stderr or '').strip()[-200:]}", "warnings": warns}))
                return
        elif fmt == "docx":
            try:
                import docx  # python-docx
            except Exception:
                print(json.dumps({"ok": False, "fmt": fmt, "error": "python-docx not installed", "warnings": warns}))
                return
            d = docx.Document()
            if title:
                d.add_heading(title, level=0)
            if author:
                d.add_paragraph(author)
            for block in md.split("\n\n"):
                b = block.strip()
                if not b:
                    continue
                hm = re.match(r"^(#{1,4})\s+(.*)$", b)
                if hm:
                    d.add_heading(hm.group(2), level=min(4, len(hm.group(1))))
                elif b.startswith(("- ", "* ")):
                    for li in b.splitlines():
                        d.add_paragraph(re.sub(r"^[-*]\s+", "", li), style="List Bullet")
                else:
                    d.add_paragraph(re.sub(r"[*_`#>]", "", b))
            d.save(out)
        else:
            print(json.dumps({"ok": False, "error": f"unknown fmt {fmt}", "warnings": warns}))
            return
        print(json.dumps({"ok": True, "out": out, "fmt": fmt, "warnings": warns}))
    except Exception as e:  # noqa
        print(json.dumps({"ok": False, "fmt": fmt, "error": f"{type(e).__name__}: {e}", "warnings": warns}))

if __name__ == "__main__":
    main()
